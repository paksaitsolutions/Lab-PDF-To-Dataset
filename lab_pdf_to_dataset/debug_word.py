from docx import Document
import re

file_path = "uploads/extracted/CBC/0000-sugra.doc"

try:
    doc = Document(file_path)
    
    print("=== PARAGRAPHS ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{i}: {para.text}")
    
    print("\n=== TABLES ===")
    print(f"Number of tables: {len(doc.tables)}")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  Row {r_idx}: {cells}")
            
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
