from docx import Document
import os
import re
import win32com.client
import pythoncom

def read_doc_with_word(doc_path):
    """Read old .doc files using MS Word COM"""
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        text = doc.Content.Text
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        return text
    except Exception as e:
        print(f"Error reading .doc file: {e}")
        return ""

def read_docx_text(docx_path):
    # Try new format first
    if docx_path.lower().endswith('.docx'):
        try:
            doc = Document(docx_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            return ""
    else:
        # Old .doc format - use Word COM
        return read_doc_with_word(os.path.abspath(docx_path))

def extract_docx_table_data(docx_path):
    """Extract structured data from Word document tables"""
    
    # For old .doc files, convert to text first
    if docx_path.lower().endswith('.doc'):
        text = read_doc_with_word(os.path.abspath(docx_path))
        # Parse text instead of tables for old format
        return parse_text_to_cbc(text)
    
    # For .docx files, use table extraction
    try:
        doc = Document(docx_path)
    except Exception as e:
        return {}
    
    # Initialize result with empty values
    result = {
        'Name': '',
        'Age': '',
        'Gender': '',
        'HB': '',
        'RBC': '',
        'HCT': '',
        'MCV': '',
        'MCH': '',
        'MCHC': '',
        'Platelets': '',
        'WBC': '',
        'Neutrophils': '',
        'Lymphocytes': '',
        'Monocytes': '',
        'Eosinophils': '',
        'ESR': ''
    }
    
    # Extract patient info from paragraphs
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # Extract Name
    name_match = re.search(r"Patient'?s?\s*Name\s*:?\s*([A-Za-z\s]+?)(?:Referring|Age|$)", full_text, re.IGNORECASE)
    if name_match:
        result['Name'] = name_match.group(1).strip()
    
    # Extract Age and Gender
    age_match = re.search(r"(\d+)\s*Year.*?(Male|Female)", full_text, re.IGNORECASE)
    if age_match:
        result['Age'] = age_match.group(1)
        result['Gender'] = age_match.group(2)
    
    # Extract test results from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue
            
            test_name = cells[0].upper().strip()
            result_value = cells[1].strip() if len(cells) > 1 else ""
            
            # Clean numeric values (remove non-numeric except decimal point)
            if result_value:
                value_match = re.search(r'([0-9]+\.?[0-9]*)', result_value)
                if value_match:
                    result_value = value_match.group(1)
            
            # Map test names
            if 'HGB' in test_name or 'HEMOGLOBIN' in test_name:
                result['HB'] = result_value
            elif 'WBC' in test_name and 'COUNT' in test_name:
                result['WBC'] = result_value
            elif 'RBC' in test_name and 'COUNT' in test_name:
                result['RBC'] = result_value
            elif 'PLATELET' in test_name:
                result['Platelets'] = result_value
            elif test_name == 'HCT' or 'HEMATOCRIT' in test_name:
                result['HCT'] = result_value
            elif test_name == 'MCV':
                result['MCV'] = result_value
            elif test_name == 'MCH' and 'MCHC' not in test_name:
                result['MCH'] = result_value
            elif test_name == 'MCHC':
                result['MCHC'] = result_value
            elif 'NEUTROPHIL' in test_name:
                result['Neutrophils'] = result_value
            elif 'LYMPHOCYTE' in test_name:
                result['Lymphocytes'] = result_value
            elif 'MONOCYTE' in test_name:
                result['Monocytes'] = result_value
            elif 'EOSINOPHIL' in test_name:
                result['Eosinophils'] = result_value
            elif test_name == 'ESR' or 'ESR' in test_name:
                result['ESR'] = result_value
    
    return result

def parse_text_to_cbc(text):
    """Parse plain text to extract CBC data"""
    result = {
        'Name': '',
        'Age': '',
        'Gender': '',
        'HB': '',
        'RBC': '',
        'HCT': '',
        'MCV': '',
        'MCH': '',
        'MCHC': '',
        'Platelets': '',
        'WBC': '',
        'Neutrophils': '',
        'Lymphocytes': '',
        'Monocytes': '',
        'Eosinophils': '',
        'ESR': ''
    }
    
    if not text:
        return result
    
    # Extract Name
    name_match = re.search(r"Patient'?s?\s*Name\s*:?\s*([A-Za-z\s]+?)(?:\r|\n|Referring|Age)", text, re.IGNORECASE)
    if name_match:
        result['Name'] = name_match.group(1).strip()
    
    # Extract Age and Gender
    age_match = re.search(r"(\d+)\s*Year.*?(Male|Female)", text, re.IGNORECASE)
    if age_match:
        result['Age'] = age_match.group(1)
        result['Gender'] = age_match.group(2)
    
    # Extract test values - look for pattern: TEST_NAME followed by VALUE (before range)
    # Pattern: Test name, then capture first number that appears (the result)
    patterns = {
        'HB': r'(?:HGB|Hemoglobin|HB)\s+([0-9]+\.?[0-9]*)\s*(?:g/dl|\s)',
        'WBC': r'WBC\s+Count\s+([0-9]+\.?[0-9]*)',
        'RBC': r'RBC\s+Count\s+([0-9]+\.?[0-9]*)',
        'Platelets': r'PLATELET\s+COUNT\s+([0-9]+)',
        'HCT': r'HCT\s+([0-9]+\.?[0-9]*)\s*%',
        'MCV': r'MCV\s+([0-9]+\.?[0-9]*)\s*(?:fL|\s)',
        'MCH': r'MCH\s+([0-9]+\.?[0-9]*)\s*(?:pg|\s)',
        'MCHC': r'MCHC\s+([0-9]+\.?[0-9]*)\s*(?:g/dL|\s)',
        'Neutrophils': r'Neutrophils\s+([0-9]+)',
        'Lymphocytes': r'Lymphocytes\s+([0-9]+)',
        'Monocytes': r'Monocytes\s+([0-9]+)',
        'Eosinophils': r'Eosinophils\s+([0-9]+)',
        'ESR': r'ESR\s+([0-9]+)'
    }
    
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result[key] = match.group(1)
    
    return result
